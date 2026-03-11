"""
LangExtract-based metadata extraction for legal documents.

Uses Google's LangExtract (Gemini-powered) to extract structured metadata
from legal documents, producing lx_results.json and interactive HTML
visualizations.

Usage:
    python -m langextract_run.lx_extract \
        --folder /path/to/docs \
        --output lx_results.json \
        --model gemini-2.5-flash
"""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv

load_dotenv()

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".doc"}

EXTRACTION_CLASSES = [
    "title",
    "document_type",
    "nuid",
    "court_name",
    "court_location",
    "judge_name",
    "filing_date",
    "plaintiff",
    "defendant",
    "clause",
]


# ── Document text extraction ─────────────────────────────────────────────

def _extract_text_pdf(path: Path) -> str:
    import fitz
    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        pages.append(page.get_text("text") or "")
    doc.close()
    return "\n".join(pages)


def _extract_text_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _extract_text_doc(path: Path) -> str:
    for cmd in ["antiword", "catdoc"]:
        try:
            result = subprocess.run(
                [cmd, str(path)],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except FileNotFoundError:
            continue
    return ""


def _extract_text_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


_LOADERS = {
    ".pdf": _extract_text_pdf,
    ".docx": _extract_text_docx,
    ".doc": _extract_text_doc,
    ".txt": _extract_text_txt,
}


def load_text(path: Path) -> str:
    ext = path.suffix.lower()
    loader = _LOADERS.get(ext)
    if not loader:
        return ""
    try:
        text = loader(path)
    except Exception as e:
        print(f"  [WARN] Failed to load {path.name}: {e}", file=sys.stderr)
        return ""
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── LangExtract examples (few-shot) ──────────────────────────────────────

def _build_examples():
    import langextract as lx

    example_text = (
        "UNITED STATES DISTRICT COURT\n"
        "SOUTHERN DISTRICT OF NEW YORK\n"
        "──────────────────\n"
        "CAMBRIA COUNTY EMPLOYEES RETIREMENT SYSTEM,\n"
        "    Plaintiff,\n"
        "v.\n"
        "GODIVA CHOCOLATIER, INC.,\n"
        "    Defendant.\n"
        "──────────────────\n"
        "Case No. 1:23-cv-00456-NRB\n"
        "COMPLAINT FOR VIOLATIONS OF FEDERAL SECURITIES LAWS\n"
        "Judge Naomi Reice Buchwald\n"
        "Filed: January 17, 2023\n"
        "First Amendment, Fifth Amendment Due Process Clause\n"
    )

    example_extractions = [
        lx.data.Extraction(
            extraction_class="title",
            extraction_text="COMPLAINT FOR VIOLATIONS OF FEDERAL SECURITIES LAWS",
            attributes={},
        ),
        lx.data.Extraction(
            extraction_class="document_type",
            extraction_text="Complaint",
            attributes={"subtype": "Securities"},
        ),
        lx.data.Extraction(
            extraction_class="nuid",
            extraction_text="1:23-cv-00456-NRB",
            attributes={},
        ),
        lx.data.Extraction(
            extraction_class="court_name",
            extraction_text="UNITED STATES DISTRICT COURT SOUTHERN DISTRICT OF NEW YORK",
            attributes={},
        ),
        lx.data.Extraction(
            extraction_class="court_location",
            extraction_text="Southern District of New York",
            attributes={},
        ),
        lx.data.Extraction(
            extraction_class="judge_name",
            extraction_text="Naomi Reice Buchwald",
            attributes={},
        ),
        lx.data.Extraction(
            extraction_class="filing_date",
            extraction_text="January 17, 2023",
            attributes={"iso": "2023-01-17"},
        ),
        lx.data.Extraction(
            extraction_class="plaintiff",
            extraction_text="CAMBRIA COUNTY EMPLOYEES RETIREMENT SYSTEM",
            attributes={},
        ),
        lx.data.Extraction(
            extraction_class="defendant",
            extraction_text="GODIVA CHOCOLATIER, INC.",
            attributes={},
        ),
        lx.data.Extraction(
            extraction_class="clause",
            extraction_text="First Amendment",
            attributes={"type": "constitutional_amendment"},
        ),
        lx.data.Extraction(
            extraction_class="clause",
            extraction_text="Fifth Amendment Due Process Clause",
            attributes={"type": "named_clause"},
        ),
    ]

    return [
        lx.data.ExampleData(
            text=example_text,
            extractions=example_extractions,
        )
    ]


PROMPT_DESCRIPTION = """\
Extract legal document metadata from court filings.

For each document, extract ALL of the following entity types that are present:

1. **title** — The document's title or heading (e.g. "Motion to Dismiss", "Complaint"). 
   Usually appears in ALL CAPS in the caption area. Do NOT use the filename.

2. **document_type** — The legal document type (e.g. Complaint, Motion, Order, 
   Memorandum of Law, Declaration, Exhibit, Stipulation, Notice, Subpoena, 
   Proof of Service, Letter, Civil Cover Sheet, Corporate Disclosure Statement).
   A document may have multiple types. Set the "subtype" attribute if applicable.

3. **nuid** — Case number / docket number (e.g. "1:23-cv-00456-NRB", "18cv339231").
   Look for patterns like "Case No.", "Docket No.", "Civil Action No.", etc.

4. **court_name** — Full name of the court (e.g. "United States District Court 
   Southern District of New York").

5. **court_location** — Geographic location of the court (e.g. "Southern District 
   of New York", "County of Santa Clara, California").

6. **judge_name** — Name of the presiding judge (e.g. "Naomi Reice Buchwald"). 
   Look for "Hon.", "Judge", "Honorable", or judge initials in docket numbers.

7. **filing_date** — The date the document was filed. Set "iso" attribute to 
   ISO-8601 format (YYYY-MM-DD). Look for "Filed:", "Date Filed:", stamps.

8. **plaintiff** — Each plaintiff or petitioner party. Extract each as a 
   separate entity. Exclude law firms, attorneys, and court personnel. Dont extract words used to refer to such parties such as "the Movant".

9. **defendant** — Each defendant or respondent party. Extract each as a 
   separate entity. Exclude law firms, attorneys, and court personnel.

10. **clause** — US Constitutional citations (amendments, named clauses, 
    articles/sections) mentioned in the document. Set "type" attribute to 
    one of: "constitutional_amendment", "named_clause", "article_section".

Focus on the caption area and first page for most metadata. Dates may appear 
anywhere. Do NOT fabricate information not present in the text.
"""


# ── Process a single document ─────────────────────────────────────────────

def extract_single(
    text: str,
    model_id: str,
    api_key: Optional[str] = None,
    examples=None,
) -> list:
    """Run LangExtract on a single document's text, return AnnotatedDocument."""
    import langextract as lx

    kwargs = dict(
        text_or_documents=text,
        prompt_description=PROMPT_DESCRIPTION,
        examples=examples,
        model_id=model_id,
        temperature=0.1,
        extraction_passes=2,
        show_progress=False,
    )
    if api_key:
        kwargs["api_key"] = api_key

    result = lx.extract(**kwargs)
    return result


# ── Convert AnnotatedDocument → result dict ───────────────────────────────

def _annotated_to_dict(ann_doc, file_path: str) -> dict:
    record = {
        "file_path": file_path,
        "title": None,
        "document_types": [],
        "nuid": None,
        "court_name": None,
        "court_location": None,
        "judge_name": None,
        "filing_date": None,
        "parties": {
            "plaintiffs": [],
            "defendants": [],
        },
        "clauses": [],
        "lx_raw_extractions": [],
    }

    for ext in ann_doc.extractions:
        cls = ext.extraction_class
        text = ext.extraction_text.replace("\n", " ").strip() if ext.extraction_text else ""
        attrs = ext.attributes or {}

        raw = {
            "class": cls,
            "text": text,
            "attributes": attrs,
        }
        if ext.char_interval:
            raw["char_start"] = ext.char_interval.start_pos
            raw["char_end"] = ext.char_interval.end_pos
        record["lx_raw_extractions"].append(raw)

        if cls == "title" and not record["title"]:
            record["title"] = text
        elif cls == "document_type":
            record["document_types"].append(text)
        elif cls == "nuid" and not record["nuid"]:
            record["nuid"] = text
        elif cls == "court_name" and not record["court_name"]:
            record["court_name"] = text
        elif cls == "court_location" and not record["court_location"]:
            record["court_location"] = text
        elif cls == "judge_name" and not record["judge_name"]:
            record["judge_name"] = text
        elif cls == "filing_date":
            iso = attrs.get("iso", text)
            if not record["filing_date"]:
                record["filing_date"] = iso
        elif cls == "plaintiff":
            if text and text not in record["parties"]["plaintiffs"]:
                record["parties"]["plaintiffs"].append(text)
        elif cls == "defendant":
            if text and text not in record["parties"]["defendants"]:
                record["parties"]["defendants"].append(text)
        elif cls == "clause":
            clause_entry = text
            if attrs.get("type"):
                clause_entry = f"{text} [{attrs['type']}]"
            if clause_entry not in record["clauses"]:
                record["clauses"].append(clause_entry)

    return record


# ── Discover files ────────────────────────────────────────────────────────

def discover_files(folder: Path) -> list[Path]:
    files = []
    for ext in SUPPORTED_EXTENSIONS:
        files.extend(folder.rglob(f"*{ext}"))
    return sorted(set(files))


# ── Main ──────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="LangExtract-based legal document metadata extraction.",
    )
    parser.add_argument(
        "--folder", required=True,
        help="Folder containing legal documents.",
    )
    parser.add_argument(
        "--output", "-o", default="langextract_run/lx_results.json",
        help="Output JSON path (default: langextract_run/lx_results.json)",
    )
    parser.add_argument(
        "--model", default="gemini-2.5-flash",
        help="LangExtract model ID (default: gemini-2.5-flash)",
    )
    parser.add_argument(
        "--api-key", default=None,
        help="API key (overrides LANGEXTRACT_API_KEY env var)",
    )
    parser.add_argument(
        "--max-docs", type=int, default=None,
        help="Process at most N documents (for testing)",
    )
    parser.add_argument(
        "--viz", action="store_true", default=True,
        help="Generate interactive HTML visualizations (default: on)",
    )
    parser.add_argument(
        "--no-viz", dest="viz", action="store_false",
        help="Skip visualization generation",
    )

    args = parser.parse_args()

    api_key = args.api_key or os.environ.get("LANGEXTRACT_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        print(
            "ERROR: No API key found. Set LANGEXTRACT_API_KEY or GOOGLE_API_KEY "
            "in .env, or pass --api-key.",
            file=sys.stderr,
        )
        sys.exit(1)

    folder = Path(args.folder)
    if not folder.is_dir():
        print(f"ERROR: '{folder}' is not a directory.", file=sys.stderr)
        sys.exit(1)

    files = discover_files(folder)
    if args.max_docs:
        files = files[:args.max_docs]

    total = len(files)
    print(f"Found {total} documents in '{folder}'", file=sys.stderr)

    examples = _build_examples()
    results = []
    annotated_docs = []
    errors = []
    total_input_chars = 0
    total_output_chars = 0

    # Monkey-patch the Gemini provider to capture real token usage from
    # the API response's usage_metadata attribute.
    _token_counts = {"prompt": 0, "candidates": 0, "total": 0, "calls": 0}

    try:
        from langextract.providers import gemini as _gemini_mod
        _orig_process = _gemini_mod.GeminiLanguageModel._process_single_prompt

        def _patched_process(self, prompt, config):
            from langextract.core import types as core_types
            from langextract import exceptions
            try:
                for key, value in self._extra_kwargs.items():
                    if key not in config and value is not None:
                        config[key] = value
                if self.gemini_schema:
                    self._validate_schema_config()
                    config.setdefault("response_mime_type", "application/json")
                    config.setdefault("response_schema", self.gemini_schema.schema_dict)

                response = self._client.models.generate_content(
                    model=self.model_id, contents=prompt, config=config
                )

                um = getattr(response, "usage_metadata", None)
                if um:
                    _token_counts["prompt"] += getattr(um, "prompt_token_count", 0) or 0
                    _token_counts["candidates"] += getattr(um, "candidates_token_count", 0) or 0
                    _token_counts["total"] += getattr(um, "total_token_count", 0) or 0
                _token_counts["calls"] += 1

                return core_types.ScoredOutput(score=1.0, output=response.text)
            except Exception as e:
                raise exceptions.InferenceRuntimeError(
                    f"Gemini API error: {str(e)}", original=e
                ) from e

        _gemini_mod.GeminiLanguageModel._process_single_prompt = _patched_process
        _patched = True
    except Exception:
        _patched = False

    try:
        from tqdm import tqdm
        file_iter = tqdm(files, desc="LangExtract", unit="doc")
    except ImportError:
        file_iter = files

    for i, fpath in enumerate(file_iter):
        rel = fpath.name
        if not hasattr(file_iter, "set_description"):
            print(f"  [{i+1}/{total}] {rel}", file=sys.stderr)

        text = load_text(fpath)
        if len(text) < 100:
            results.append({
                "file_path": str(fpath),
                "skipped": True,
                "skip_reason": "too_little_text",
            })
            continue

        try:
            ann = extract_single(
                text=text,
                model_id=args.model,
                api_key=api_key,
                examples=examples,
            )

            if isinstance(ann, list):
                ann_doc = ann[0] if ann else None
            else:
                ann_doc = ann

            if ann_doc is None:
                results.append({
                    "file_path": str(fpath),
                    "skipped": True,
                    "skip_reason": "no_extraction_result",
                })
                continue

            record = _annotated_to_dict(ann_doc, str(fpath))
            results.append(record)
            annotated_docs.append((ann_doc, str(fpath)))

            total_input_chars += len(text)
            out_chars = sum(
                len(e.extraction_text or "")
                for e in ann_doc.extractions
            )
            total_output_chars += out_chars

        except Exception as e:
            err_msg = f"{type(e).__name__}: {e}"
            print(f"  [ERROR] {rel}: {err_msg}", file=sys.stderr)
            errors.append({"file": str(fpath), "error": err_msg})
            results.append({
                "file_path": str(fpath),
                "skipped": True,
                "skip_reason": f"error: {err_msg}",
            })

    # Restore original method if patched
    if _patched:
        _gemini_mod.GeminiLanguageModel._process_single_prompt = _orig_process

    # ── Write results ─────────────────────────────────────────────────────
    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print(f"\nResults saved to {out_path} ({len(results)} documents)", file=sys.stderr)

    # ── Summary stats ─────────────────────────────────────────────────────
    processed = [r for r in results if not r.get("skipped")]
    skipped = [r for r in results if r.get("skipped")]

    stats = {
        "total_documents": total,
        "processed": len(processed),
        "skipped": len(skipped),
        "errors": len(errors),
        "fields": {},
    }

    for field_name in ["title", "nuid", "court_name", "court_location",
                       "judge_name", "filing_date"]:
        filled = sum(1 for r in processed if r.get(field_name))
        stats["fields"][field_name] = {
            "filled": filled,
            "total": len(processed),
            "pct": round(100 * filled / max(len(processed), 1), 1),
        }

    has_plaintiffs = sum(1 for r in processed if r.get("parties", {}).get("plaintiffs"))
    has_defendants = sum(1 for r in processed if r.get("parties", {}).get("defendants"))
    has_doctypes = sum(1 for r in processed if r.get("document_types"))
    has_clauses = sum(1 for r in processed if r.get("clauses"))

    stats["fields"]["plaintiffs"] = {
        "filled": has_plaintiffs,
        "total": len(processed),
        "pct": round(100 * has_plaintiffs / max(len(processed), 1), 1),
    }
    stats["fields"]["defendants"] = {
        "filled": has_defendants,
        "total": len(processed),
        "pct": round(100 * has_defendants / max(len(processed), 1), 1),
    }
    stats["fields"]["document_types"] = {
        "filled": has_doctypes,
        "total": len(processed),
        "pct": round(100 * has_doctypes / max(len(processed), 1), 1),
    }
    stats["fields"]["clauses"] = {
        "filled": has_clauses,
        "total": len(processed),
        "pct": round(100 * has_clauses / max(len(processed), 1), 1),
    }

    if _patched and _token_counts["calls"] > 0:
        stats["token_usage"] = {
            "source": "gemini_api (usage_metadata)",
            "api_calls": _token_counts["calls"],
            "input_tokens": _token_counts["prompt"],
            "output_tokens": _token_counts["candidates"],
            "total_tokens": _token_counts["total"],
            "total_input_chars": total_input_chars,
            "total_output_chars": total_output_chars,
        }
    else:
        est_input_tokens = total_input_chars // 4
        est_output_tokens = total_output_chars // 4
        est_input_tokens_with_passes = est_input_tokens * 2
        stats["token_usage"] = {
            "source": "estimated (~4 chars/token)",
            "api_calls": 0,
            "input_tokens": est_input_tokens_with_passes,
            "output_tokens": est_output_tokens,
            "total_tokens": est_input_tokens_with_passes + est_output_tokens,
            "total_input_chars": total_input_chars,
            "total_output_chars": total_output_chars,
        }

    stats_path = out_path.parent / "lx_stats.json"
    with open(stats_path, "w", encoding="utf-8") as f:
        json.dump(stats, f, indent=2)
    print(f"Stats saved to {stats_path}", file=sys.stderr)

    _print_summary_table(stats)

    # ── Visualizations ────────────────────────────────────────────────────
    if args.viz and annotated_docs:
        _generate_visualizations(annotated_docs, out_path, stats)

    if errors:
        err_path = out_path.parent / "lx_errors.json"
        with open(err_path, "w", encoding="utf-8") as f:
            json.dump(errors, f, indent=2)
        print(f"Errors logged to {err_path}", file=sys.stderr)


def _print_summary_table(stats: dict):
    print("\n" + "=" * 60, file=sys.stderr)
    print("  LangExtract Extraction Summary", file=sys.stderr)
    print("=" * 60, file=sys.stderr)
    print(f"  Total documents:  {stats['total_documents']}", file=sys.stderr)
    print(f"  Processed:        {stats['processed']}", file=sys.stderr)
    print(f"  Skipped:          {stats['skipped']}", file=sys.stderr)
    print(f"  Errors:           {stats['errors']}", file=sys.stderr)
    print("-" * 60, file=sys.stderr)
    print(f"  {'Field':<20} {'Filled':>8} {'Total':>8} {'%':>8}", file=sys.stderr)
    print("-" * 60, file=sys.stderr)
    for field_name, info in stats["fields"].items():
        print(
            f"  {field_name:<20} {info['filled']:>8} {info['total']:>8} {info['pct']:>7.1f}%",
            file=sys.stderr,
        )
    print("-" * 60, file=sys.stderr)

    tu = stats.get("token_usage", {})
    if tu:
        source = tu.get("source", "unknown")
        print(f"  Token Usage ({source})", file=sys.stderr)
        print("-" * 60, file=sys.stderr)
        print(f"  API calls:            {tu.get('api_calls', 0):>12,}", file=sys.stderr)
        print(f"  Input tokens:         {tu.get('input_tokens', 0):>12,}", file=sys.stderr)
        print(f"  Output tokens:        {tu.get('output_tokens', 0):>12,}", file=sys.stderr)
        print(f"  Total tokens:         {tu.get('total_tokens', 0):>12,}", file=sys.stderr)
        print(f"  Input chars:          {tu.get('total_input_chars', 0):>12,}", file=sys.stderr)
        print(f"  Output chars:         {tu.get('total_output_chars', 0):>12,}", file=sys.stderr)

    print("=" * 60 + "\n", file=sys.stderr)


_LX_FIELD_COLORS_RGB: dict[str, tuple[float, float, float]] = {
    "title":          (1.0, 0.84, 0.0),
    "document_type":  (0.53, 0.81, 0.92),
    "nuid":           (0.60, 0.98, 0.60),
    "court_name":     (0.87, 0.63, 0.87),
    "court_location": (0.94, 0.63, 0.63),
    "judge_name":     (1.0, 0.63, 0.48),
    "filing_date":    (0.56, 0.93, 0.56),
    "plaintiff":      (1.0, 0.71, 0.76),
    "defendant":      (1.0, 0.71, 0.76),
    "clause":         (1.0, 0.39, 0.28),
}


def _annotate_pdf_lx(ann_doc, src_path: str, out_path: Path) -> bool:
    """Annotate a PDF with LangExtract extraction highlights."""
    import fitz

    if not src_path.lower().endswith(".pdf"):
        return False
    try:
        doc = fitz.open(src_path)
    except Exception:
        return False

    n_pages = len(doc)
    highlights = 0

    for ext in ann_doc.extractions:
        cls = ext.extraction_class or ""
        span = (ext.extraction_text or "").strip()
        if not span:
            continue

        color = _LX_FIELD_COLORS_RGB.get(cls, (0.8, 0.8, 0.8))

        # Try every page — LangExtract doesn't always give page numbers
        found = False
        for page in doc:
            quads = page.search_for(span[:60], quads=True)
            if quads:
                annot = page.add_highlight_annot(quads)
                annot.set_colors(stroke=color)
                annot.set_opacity(0.45)
                info = annot.info
                info["content"] = f"[{cls}] {span[:120]}"
                info["title"] = cls
                annot.set_info(info)
                annot.update()
                highlights += 1
                found = True
                break
        if not found and len(span) > 20:
            for page in doc:
                quads = page.search_for(span[:20], quads=True)
                if quads:
                    annot = page.add_highlight_annot(quads)
                    annot.set_colors(stroke=color)
                    annot.set_opacity(0.45)
                    info = annot.info
                    info["content"] = f"[{cls}] {span[:120]}"
                    info["title"] = cls
                    annot.set_info(info)
                    annot.update()
                    highlights += 1
                    break

    if highlights == 0:
        doc.close()
        return False

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path), garbage=3, deflate=True)
    doc.close()
    return True


def _generate_visualizations(annotated_docs, out_path: Path, stats: dict):
    import langextract as lx

    viz_dir = out_path.parent / "visualizations"
    viz_dir.mkdir(exist_ok=True)

    # Save annotated JSONL
    raw_docs = [pair[0] if isinstance(pair, tuple) else pair for pair in annotated_docs]
    jsonl_path = viz_dir / "lx_annotated.jsonl"
    try:
        lx.io.save_annotated_documents(raw_docs, str(jsonl_path))
        print(f"Annotated JSONL saved to {jsonl_path}", file=sys.stderr)
    except Exception as e:
        print(f"  [WARN] Could not save JSONL: {e}", file=sys.stderr)

    # Generate annotated PDFs
    sample = annotated_docs[:20]
    pdf_count = 0
    for idx, pair in enumerate(sample):
        if isinstance(pair, tuple):
            ann_doc, fpath = pair
        else:
            continue

        safe_name = Path(fpath).stem[:60].replace(" ", "_")
        pdf_path = viz_dir / f"doc_{idx:03d}_{safe_name}.pdf"
        try:
            ok = _annotate_pdf_lx(ann_doc, fpath, pdf_path)
            if ok:
                pdf_count += 1
        except Exception as e:
            print(f"  [WARN] PDF annotation failed for doc {idx}: {e}", file=sys.stderr)

    if pdf_count:
        print(
            f"Annotated PDFs saved to {viz_dir}/ ({pdf_count} docs)",
            file=sys.stderr,
        )

    try:
        _generate_charts(stats, viz_dir)
    except Exception as e:
        print(f"  [WARN] Chart generation failed: {e}", file=sys.stderr)


def _generate_charts(stats: dict, viz_dir: Path):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.ticker as mticker

    fields = list(stats["fields"].keys())
    pcts = [stats["fields"][f]["pct"] for f in fields]
    filled = [stats["fields"][f]["filled"] for f in fields]
    totals = [stats["fields"][f]["total"] for f in fields]

    fig, axes = plt.subplots(1, 2, figsize=(16, 6))
    fig.suptitle("LangExtract — Legal Document Metadata Extraction", fontsize=14, fontweight="bold")

    ax1 = axes[0]
    colors = []
    for p in pcts:
        if p >= 80:
            colors.append("#2ecc71")
        elif p >= 50:
            colors.append("#f39c12")
        else:
            colors.append("#e74c3c")

    bars1 = ax1.barh(fields, pcts, color=colors, edgecolor="white", height=0.6)
    ax1.set_xlabel("Coverage (%)")
    ax1.set_title("Field Coverage Rate")
    ax1.set_xlim(0, 105)
    ax1.xaxis.set_major_formatter(mticker.PercentFormatter())
    for bar, pct in zip(bars1, pcts):
        ax1.text(bar.get_width() + 1, bar.get_y() + bar.get_height() / 2,
                 f"{pct:.1f}%", va="center", fontsize=9)

    ax2 = axes[1]
    ax2.barh(fields, totals, color="#dfe6e9", edgecolor="white", height=0.6, label="Total")
    ax2.barh(fields, filled, color="#0984e3", edgecolor="white", height=0.6, label="Filled")
    ax2.set_xlabel("Document Count")
    ax2.set_title("Filled vs Total Documents")
    ax2.legend(loc="lower right")

    plt.tight_layout()
    chart_path = viz_dir / "extraction_summary.png"
    fig.savefig(str(chart_path), dpi=150, bbox_inches="tight")
    plt.close(fig)
    print(f"Summary chart saved to {chart_path}", file=sys.stderr)

    fig2, ax3 = plt.subplots(figsize=(8, 8))
    non_zero = [(f, v) for f, v in zip(fields, filled) if v > 0]
    if non_zero:
        labels, values = zip(*non_zero)
        wedges, texts, autotexts = ax3.pie(
            values, labels=labels, autopct="%1.1f%%",
            colors=plt.cm.Set3.colors[:len(values)],
            startangle=140,
        )
        ax3.set_title("Extraction Class Distribution (by filled count)")
        pie_path = viz_dir / "extraction_distribution.png"
        fig2.savefig(str(pie_path), dpi=150, bbox_inches="tight")
        print(f"Distribution chart saved to {pie_path}", file=sys.stderr)
    plt.close(fig2)


if __name__ == "__main__":
    main()
