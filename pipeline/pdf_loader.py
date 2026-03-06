"""
Document loading for the legal document pipeline.

Supported formats:
  .pdf   — PyMuPDF (fitz)
  .docx  — python-docx
  .doc   — antiword or catdoc subprocess
  .txt   — plain UTF-8 read

All loaders return a LoadedDocument with a unified interface.
"""

from __future__ import annotations

import re
import subprocess
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from pipeline.config import MIN_CHARS

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".doc"}

# Characters per synthetic "page" for non-PDF formats (plain text / DOCX).
# Legal pages are typically 2 500–3 500 chars; 3 000 is a good midpoint.
_CHARS_PER_PAGE = 3_000


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class PageInfo:
    page_num: int     # 1-indexed
    text: str
    char_count: int


@dataclass
class LoadedDocument:
    pdf_path: str
    full_text: str
    pages: list[PageInfo]
    metadata: dict
    total_chars: int
    skipped: bool
    skip_reason: str
    first_page_text: str = ""
    page_blocks: list[dict] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

_EM_DASH  = re.compile(r"\u2014|\u2013|\u2012")
_LIGATURE = re.compile(r"\ufb00|\ufb01|\ufb02|\ufb03|\ufb04")
_LIGATURE_MAP = {"\ufb00": "ff", "\ufb01": "fi", "\ufb02": "fl",
                 "\ufb03": "ffi", "\ufb04": "ffl"}


def _normalize_dashes(text: str) -> str:
    text = _EM_DASH.sub("-", text)
    for lig, rep in _LIGATURE_MAP.items():
        text = text.replace(lig, rep)
    return text


def _text_to_pages(text: str) -> list[PageInfo]:
    """
    Split a flat string into synthetic PageInfo objects (~_CHARS_PER_PAGE chars each).
    Breaks on newlines so lines are never split mid-way.
    """
    lines = text.splitlines(keepends=True)
    pages: list[PageInfo] = []
    page_num = 1
    buf: list[str] = []
    buf_len = 0

    for line in lines:
        buf.append(line)
        buf_len += len(line)
        if buf_len >= _CHARS_PER_PAGE:
            page_text = "".join(buf)
            pages.append(PageInfo(page_num=page_num,
                                  text=page_text,
                                  char_count=len(page_text)))
            page_num += 1
            buf = []
            buf_len = 0

    if buf:
        page_text = "".join(buf)
        pages.append(PageInfo(page_num=page_num,
                              text=page_text,
                              char_count=len(page_text)))

    return pages or [PageInfo(page_num=1, text="", char_count=0)]


def _wrap_plain_text(
    path: Path,
    full_text: str,
    metadata: dict,
    min_chars: int,
) -> LoadedDocument:
    """Wrap extracted plain text into a LoadedDocument."""
    total_chars = len(full_text.strip())

    if total_chars < min_chars:
        return LoadedDocument(
            pdf_path=str(path),
            full_text="",
            pages=[],
            metadata=metadata,
            total_chars=total_chars,
            skipped=True,
            skip_reason="too_little_text",
        )

    pages = _text_to_pages(full_text)
    first_page_text = pages[0].text if pages else ""

    return LoadedDocument(
        pdf_path=str(path),
        full_text=full_text,
        pages=pages,
        metadata=metadata,
        total_chars=total_chars,
        skipped=False,
        skip_reason="",
        first_page_text=first_page_text,
        page_blocks=[],
    )


# ---------------------------------------------------------------------------
# PDF loader (PyMuPDF)
# ---------------------------------------------------------------------------

def _extract_first_page_zones(page) -> list[dict]:
    """Extract text blocks from the first PDF page for layout analysis."""
    try:
        blocks = page.get_text("dict", flags=0).get("blocks", [])
        zones = []
        for b in blocks:
            if b.get("type") != 0:   # 0 = text block
                continue
            lines_text = []
            for line in b.get("lines", []):
                span_text = " ".join(s.get("text", "") for s in line.get("spans", []))
                if span_text.strip():
                    lines_text.append(span_text)
            if lines_text:
                zones.append({
                    "bbox": b.get("bbox"),
                    "text": "\n".join(lines_text),
                })
        return zones
    except Exception:
        return []


def _read_pdf(
    path: Path,
    min_chars: int,
    extract_first_page_blocks: bool,
) -> LoadedDocument:
    """Load a PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "PyMuPDF is required for PDF support. "
            "Install with: pip install PyMuPDF"
        ) from exc

    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        return LoadedDocument(
            pdf_path=str(path),
            full_text="",
            pages=[],
            metadata={"format": "pdf", "filename": path.name},
            total_chars=0,
            skipped=True,
            skip_reason=f"pdf_open_error:{exc}",
        )

    pages: list[PageInfo] = []
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text") or ""
        text = _normalize_dashes(text)
        pages.append(PageInfo(page_num=page_num,
                              text=text,
                              char_count=len(text)))

    doc.close()

    full_text = "\n".join(p.text for p in pages)
    total_chars = len(full_text.strip())

    if total_chars < min_chars:
        return LoadedDocument(
            pdf_path=str(path),
            full_text="",
            pages=[],
            metadata={"format": "pdf", "filename": path.name},
            total_chars=total_chars,
            skipped=True,
            skip_reason="too_little_text",
        )

    first_page_text = pages[0].text if pages else ""
    page_blocks: list[dict] = []

    if extract_first_page_blocks and pages:
        try:
            doc2 = fitz.open(str(path))
            page_blocks = _extract_first_page_zones(doc2[0])
            doc2.close()
        except Exception:
            page_blocks = []

    return LoadedDocument(
        pdf_path=str(path),
        full_text=full_text,
        pages=pages,
        metadata={"format": "pdf", "filename": path.name},
        total_chars=total_chars,
        skipped=False,
        skip_reason="",
        first_page_text=first_page_text,
        page_blocks=page_blocks,
    )


# ---------------------------------------------------------------------------
# DOCX loader (python-docx)
# ---------------------------------------------------------------------------

def _read_docx(path: Path) -> tuple[str, dict]:
    """
    Read a .docx file using python-docx.
    Iterates body children in document order so that table cells embedded in
    the case caption (e.g. judge name / case number) appear at the correct
    position rather than being appended after all paragraphs.
    """
    try:
        from docx import Document as DocxDocument  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for .docx support. "
            "Install with: pip install python-docx"
        ) from exc

    doc = DocxDocument(str(path))
    parts: list[str] = []

    from docx.text.paragraph import Paragraph as DocxParagraph  # type: ignore
    from docx.table import Table as DocxTable                   # type: ignore

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
        if tag == "p":
            para = DocxParagraph(block, doc)
            if para.text.strip():
                parts.append(para.text)
        elif tag == "tbl":
            table = DocxTable(block, doc)
            for row in table.rows:
                seen: set = set()
                row_cells: list[str] = []
                for cell in row.cells:
                    if id(cell._tc) not in seen and cell.text.strip():
                        seen.add(id(cell._tc))
                        row_cells.append(cell.text.strip())
                if row_cells:
                    # Join multi-column rows so fields like "Hon. X" and
                    # "Case No." are still extractable from the same line.
                    parts.append("  |  ".join(row_cells))

    cp = doc.core_properties
    metadata = {
        "format":   "docx",
        "filename": path.name,
        "author":   cp.author or "",
        "title":    cp.title or "",
        "created":  str(cp.created) if cp.created else "",
        "modified": str(cp.modified) if cp.modified else "",
    }

    return _normalize_dashes("\n".join(parts)), metadata


# ---------------------------------------------------------------------------
# DOC loader (legacy — antiword / catdoc)
# ---------------------------------------------------------------------------

def _read_doc(path: Path) -> tuple[str, dict]:
    """
    Read a legacy .doc file via antiword or catdoc subprocess.
    Raises RuntimeError if neither tool is available.
    """
    metadata = {"format": "doc", "filename": path.name}

    for tool in ("antiword", "catdoc"):
        try:
            result = subprocess.run(
                [tool, str(path)],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return _normalize_dashes(result.stdout), metadata
        except FileNotFoundError:
            continue
        except subprocess.TimeoutExpired:
            break

    raise RuntimeError(
        f"Cannot read '{path.name}': neither antiword nor catdoc is available. "
        "Install one with: sudo apt install antiword"
    )


# ---------------------------------------------------------------------------
# TXT loader
# ---------------------------------------------------------------------------

def _read_txt(path: Path) -> tuple[str, dict]:
    text = path.read_text(errors="replace")
    return _normalize_dashes(text), {"format": "txt", "filename": path.name}


# ---------------------------------------------------------------------------
# Path detection
# ---------------------------------------------------------------------------

def _looks_like_path(source: str | Path) -> bool:
    """Return True if source looks like a filesystem path (no newlines, ≤4096 chars)."""
    s = str(source)
    if "\n" in s or len(s) > 4096:
        return False
    try:
        Path(s)
        return True
    except (OSError, ValueError):
        return False


# ---------------------------------------------------------------------------
# Unified dispatcher
# ---------------------------------------------------------------------------

def load_document(
    source: str | Path,
    min_chars: int = MIN_CHARS,
    extract_first_page_blocks: bool = True,
) -> LoadedDocument:
    """
    Load any supported document type and return a LoadedDocument.

    Parameters
    ----------
    source                   : file path (str or Path) or raw text string
    min_chars                : minimum characters; shorter docs are skipped
    extract_first_page_blocks: PDF only — extract first-page layout blocks
    """
    # Inline text string (not a path)
    if not isinstance(source, Path) and not _looks_like_path(source):
        raw = str(source)
        return _wrap_plain_text(
            Path("<inline>"), _normalize_dashes(raw),
            {"format": "text", "filename": "<inline>"}, min_chars,
        )

    path = Path(source)

    if not path.exists():
        return LoadedDocument(
            pdf_path=str(path),
            full_text="",
            pages=[],
            metadata={"format": "unknown", "filename": path.name},
            total_chars=0,
            skipped=True,
            skip_reason="file_not_found",
        )

    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            return _read_pdf(path, min_chars, extract_first_page_blocks)

        elif suffix == ".docx":
            raw, meta = _read_docx(path)
            return _wrap_plain_text(path, raw, meta, min_chars)

        elif suffix == ".doc":
            raw, meta = _read_doc(path)
            return _wrap_plain_text(path, raw, meta, min_chars)

        elif suffix == ".txt":
            raw, meta = _read_txt(path)
            return _wrap_plain_text(path, raw, meta, min_chars)

        else:
            return LoadedDocument(
                pdf_path=str(path),
                full_text="",
                pages=[],
                metadata={"format": "unsupported", "filename": path.name},
                total_chars=0,
                skipped=True,
                skip_reason=f"unsupported_extension:{suffix}",
            )

    except Exception as exc:
        return LoadedDocument(
            pdf_path=str(path),
            full_text="",
            pages=[],
            metadata={"format": "error", "filename": path.name},
            total_chars=0,
            skipped=True,
            skip_reason=f"load_error:{exc}",
        )
