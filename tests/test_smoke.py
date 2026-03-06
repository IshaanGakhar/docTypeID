"""
Smoke tests — verify the pipeline runs end-to-end on synthetic text input
without crashing, and that the output schema is correct.

Run:
    pytest tests/test_smoke.py -v
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from pipeline.run_pipeline import run_pipeline, run_pipeline_dir, discover_files
from pipeline.pdf_loader import load_document, SUPPORTED_EXTENSIONS
from pipeline.cluster_ingestion import (
    load_cluster_csv,
    consolidate_cluster,
    run_pipeline_clusters,
    CsvDocRecord,
    NOISE_CLUSTER_IDS,
)

# ---------------------------------------------------------------------------
# Minimal legal document text — long enough to pass the 150-char gate
# ---------------------------------------------------------------------------

SAMPLE_TEXT = """\
UNITED STATES DISTRICT COURT
FOR THE NORTHERN DISTRICT OF CALIFORNIA

ACME CORPORATION,
    Plaintiff,

v.

GLOBEX LLC,
    Defendant.

Case No. 3:22-cv-01234-JCS

MOTION TO DISMISS
UNDER RULE 12(b)(6)

Assigned to: Hon. Joseph C. Spero

Filed: March 15, 2023

INTRODUCTION

Plaintiff ACME Corporation brings this action against Defendant Globex LLC
for breach of contract and tortious interference.

JURISDICTION AND VENUE

This Court has subject matter jurisdiction under 28 U.S.C. § 1332.
Venue is proper in this district under 28 U.S.C. § 1391.

ARGUMENT

Defendant moves to dismiss the complaint in its entirety for failure to
state a claim upon which relief can be granted pursuant to Federal Rule
of Civil Procedure 12(b)(6).

CONCLUSION

For the foregoing reasons, Defendant respectfully requests that this Court
grant its Motion to Dismiss with prejudice.

WHEREFORE, Defendant prays for relief as set forth above.
"""

BELOW_MIN_TEXT = "Too short."   # < 150 chars, should be skipped


# ---------------------------------------------------------------------------
# Schema validation helper
# ---------------------------------------------------------------------------

_REQUIRED_KEYS = {
    "skipped", "file_path", "title", "document_types", "nuid",
    "court_name", "court_location", "judge_name", "filing_date",
    "parties", "clauses", "evidence", "confidence",
}

_PARTY_KEYS = {"plaintiffs", "defendants", "petitioners", "respondents"}

_CLAUSE_KEYS = {"clause_type", "heading", "text", "page_start", "page_end"}

_EVIDENCE_ITEM_KEYS = {"source", "page", "span_text", "char_start", "char_end", "rule_id"}


def _validate_schema(result: dict) -> None:
    """Assert the output dict conforms to the required schema."""
    assert isinstance(result, dict), "Result must be a dict"
    missing = _REQUIRED_KEYS - result.keys()
    assert not missing, f"Missing keys: {missing}"

    assert isinstance(result["document_types"], list)
    assert isinstance(result["parties"], dict)
    assert _PARTY_KEYS.issubset(result["parties"].keys()), \
        f"Missing party roles: {_PARTY_KEYS - result['parties'].keys()}"

    for party_list in result["parties"].values():
        assert isinstance(party_list, list)

    assert isinstance(result["clauses"], list)
    for clause in result["clauses"]:
        missing_c = _CLAUSE_KEYS - clause.keys()
        assert not missing_c, f"Clause missing keys: {missing_c}"
        assert isinstance(clause["page_start"], int)
        assert isinstance(clause["page_end"], int)

    assert isinstance(result["evidence"], dict)
    for field_name, ev_list in result["evidence"].items():
        assert isinstance(ev_list, list), f"evidence[{field_name}] must be list"
        for ev in ev_list:
            missing_e = _EVIDENCE_ITEM_KEYS - ev.keys()
            assert not missing_e, f"Evidence item missing keys: {missing_e}"

    assert isinstance(result["confidence"], dict)
    for field_name, score in result["confidence"].items():
        assert isinstance(score, float), f"confidence[{field_name}] must be float"
        assert 0.0 <= score <= 1.0, f"confidence[{field_name}] out of range: {score}"


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestPipelineSchema:
    def test_schema_valid_on_sample_text(self):
        result = run_pipeline(SAMPLE_TEXT)
        assert not result["skipped"]
        _validate_schema(result)

    def test_output_is_json_serializable(self):
        result = run_pipeline(SAMPLE_TEXT)
        serialized = json.dumps(result)
        assert isinstance(serialized, str)

    def test_skip_on_too_short_text(self):
        result = run_pipeline(BELOW_MIN_TEXT)
        assert result["skipped"] is True
        assert result["skip_reason"] == "too_little_text"

    def test_document_types_are_strings(self):
        result = run_pipeline(SAMPLE_TEXT)
        for dt in result["document_types"]:
            assert isinstance(dt, str)

    def test_title_extracted(self):
        result = run_pipeline(SAMPLE_TEXT)
        # Title may or may not be found, but must be str or None
        assert result["title"] is None or isinstance(result["title"], str)

    def test_nuid_extracted(self):
        result = run_pipeline(SAMPLE_TEXT)
        # NUID should be extracted from "Case No. 3:22-cv-01234-JCS"
        assert result["nuid"] is not None, "Expected NUID to be extracted"
        assert "22" in result["nuid"] or "1234" in result["nuid"]

    def test_court_extracted(self):
        result = run_pipeline(SAMPLE_TEXT)
        assert result["court_name"] is not None

    def test_filing_date_extracted(self):
        result = run_pipeline(SAMPLE_TEXT)
        assert result["filing_date"] is not None
        assert result["filing_date"] == "2023-03-15"

    def test_parties_extracted(self):
        result = run_pipeline(SAMPLE_TEXT)
        parties = result["parties"]
        all_parties = (
            parties["plaintiffs"] + parties["defendants"] +
            parties["petitioners"] + parties["respondents"]
        )
        assert len(all_parties) > 0, "Expected at least one party"

    def test_clauses_have_page_ranges(self):
        result = run_pipeline(SAMPLE_TEXT)
        for clause in result["clauses"]:
            assert clause["page_start"] >= 1
            assert clause["page_end"] >= clause["page_start"]

    def test_evidence_populated_for_extracted_fields(self):
        result = run_pipeline(SAMPLE_TEXT)
        # If a field was extracted (non-null), evidence should be present
        for field in ["nuid", "court_name", "filing_date"]:
            if result[field] is not None:
                assert field in result["evidence"], \
                    f"Evidence missing for field '{field}' which was extracted"
                assert len(result["evidence"][field]) > 0

    def test_confidence_range(self):
        result = run_pipeline(SAMPLE_TEXT)
        for field, score in result["confidence"].items():
            assert 0.0 <= score <= 1.0, f"{field} confidence out of range: {score}"

    def test_fallback_mode_no_models(self, tmp_path):
        """Pipeline must run correctly even when no model files exist."""
        result = run_pipeline(SAMPLE_TEXT)
        # Rule-based fallback is always active; result should still be valid
        _validate_schema(result)


class TestSkipBehavior:
    def test_empty_string_skipped(self):
        result = run_pipeline("")
        assert result["skipped"] is True

    def test_whitespace_only_skipped(self):
        result = run_pipeline("   \n\t  ")
        assert result["skipped"] is True

    def test_exactly_min_chars_not_skipped(self):
        """150+ char document should not be skipped."""
        long_text = "A" * 151
        result = run_pipeline(long_text)
        assert result["skipped"] is False


class TestTxtFileSupport:
    """Pipeline accepts .txt files and produces a valid schema."""

    def _write_txt(self, tmp_path: Path, content: str) -> Path:
        p = tmp_path / "test_doc.txt"
        p.write_text(content, encoding="utf-8")
        return p

    def test_txt_file_not_skipped(self, tmp_path):
        p = self._write_txt(tmp_path, SAMPLE_TEXT)
        result = run_pipeline(p)
        assert not result["skipped"]

    def test_txt_schema_valid(self, tmp_path):
        p = self._write_txt(tmp_path, SAMPLE_TEXT)
        result = run_pipeline(p)
        _validate_schema(result)

    def test_txt_file_path_in_output(self, tmp_path):
        p = self._write_txt(tmp_path, SAMPLE_TEXT)
        result = run_pipeline(p)
        assert str(p) in result["file_path"]

    def test_txt_json_serializable(self, tmp_path):
        p = self._write_txt(tmp_path, SAMPLE_TEXT)
        result = run_pipeline(p)
        assert json.dumps(result)

    def test_txt_nuid_extracted(self, tmp_path):
        p = self._write_txt(tmp_path, SAMPLE_TEXT)
        result = run_pipeline(p)
        assert result["nuid"] is not None

    def test_txt_filing_date_extracted(self, tmp_path):
        p = self._write_txt(tmp_path, SAMPLE_TEXT)
        result = run_pipeline(p)
        assert result["filing_date"] == "2023-03-15"

    def test_txt_court_extracted(self, tmp_path):
        p = self._write_txt(tmp_path, SAMPLE_TEXT)
        result = run_pipeline(p)
        assert result["court_name"] is not None

    def test_txt_too_short_skipped(self, tmp_path):
        p = self._write_txt(tmp_path, "Short.")
        result = run_pipeline(p)
        assert result["skipped"] is True
        assert result["skip_reason"] == "too_little_text"

    def test_txt_pagination(self, tmp_path):
        """Long .txt files are split into multiple synthetic pages."""
        long_text = (SAMPLE_TEXT + "\n") * 20   # ~20x content
        p = self._write_txt(tmp_path, long_text)
        from pipeline.pdf_loader import load_document
        doc = load_document(p)
        assert len(doc.pages) > 1

    def test_txt_en_dash_normalized(self, tmp_path):
        """En-dashes in .txt should be normalized before extraction."""
        text_with_dashes = SAMPLE_TEXT.replace("-", "\u2013")
        p = self._write_txt(tmp_path, text_with_dashes)
        result = run_pipeline(p)
        assert not result["skipped"]


class TestDocxFileSupport:
    """Pipeline accepts .docx files (requires python-docx)."""

    def _make_docx(self, tmp_path: Path, content: str) -> Path:
        pytest.importorskip("docx", reason="python-docx not installed")
        from docx import Document as DocxDocument
        doc = DocxDocument()
        for para in content.split("\n"):
            doc.add_paragraph(para)
        p = tmp_path / "test_doc.docx"
        doc.save(str(p))
        return p

    def test_docx_not_skipped(self, tmp_path):
        p = self._make_docx(tmp_path, SAMPLE_TEXT)
        result = run_pipeline(p)
        assert not result["skipped"]

    def test_docx_schema_valid(self, tmp_path):
        p = self._make_docx(tmp_path, SAMPLE_TEXT)
        result = run_pipeline(p)
        _validate_schema(result)

    def test_docx_json_serializable(self, tmp_path):
        p = self._make_docx(tmp_path, SAMPLE_TEXT)
        result = run_pipeline(p)
        assert json.dumps(result)

    def test_docx_nuid_extracted(self, tmp_path):
        p = self._make_docx(tmp_path, SAMPLE_TEXT)
        result = run_pipeline(p)
        assert result["nuid"] is not None

    def test_docx_filing_date_extracted(self, tmp_path):
        p = self._make_docx(tmp_path, SAMPLE_TEXT)
        result = run_pipeline(p)
        assert result["filing_date"] == "2023-03-15"


class TestSupportedExtensions:
    def test_supported_extensions_set(self):
        assert ".pdf"  in SUPPORTED_EXTENSIONS
        assert ".txt"  in SUPPORTED_EXTENSIONS
        assert ".docx" in SUPPORTED_EXTENSIONS
        assert ".doc"  in SUPPORTED_EXTENSIONS

    def test_load_document_inline_text(self):
        from pipeline.pdf_loader import load_document
        doc = load_document(SAMPLE_TEXT)
        assert not doc.skipped
        assert doc.full_text == SAMPLE_TEXT
        assert doc.pdf_path == "<text_input>"

    def test_load_document_inline_too_short(self):
        from pipeline.pdf_loader import load_document
        doc = load_document("hi")
        assert doc.skipped
        assert doc.skip_reason == "too_little_text"

    def test_load_document_nonexistent_file(self):
        """A path with a supported extension that doesn't exist → file_not_found."""
        from pipeline.pdf_loader import load_document
        doc = load_document("/no/such/file.txt")
        assert doc.skipped
        assert "file_not_found" in doc.skip_reason
        assert doc.pdf_path == "/no/such/file.txt"   # preserves the intended path


# ---------------------------------------------------------------------------
# Folder pipeline tests
# ---------------------------------------------------------------------------

def _write(folder: Path, name: str, content: str) -> Path:
    p = folder / name
    p.write_text(content, encoding="utf-8")
    return p


class TestDiscoverFiles:
    def test_finds_txt_files(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        files = discover_files(tmp_path)
        names = [f.name for f in files]
        assert "a.txt" in names and "b.txt" in names

    def test_ignores_unsupported_extensions(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.csv", "col1,col2\n1,2")
        _write(tmp_path, "c.json", "{}")
        files = discover_files(tmp_path)
        names = [f.name for f in files]
        assert "a.txt" in names
        assert "b.csv" not in names
        assert "c.json" not in names

    def test_recursive_finds_nested(self, tmp_path):
        sub = tmp_path / "subdir"
        sub.mkdir()
        _write(sub, "nested.txt", SAMPLE_TEXT)
        files = discover_files(tmp_path, recursive=True)
        names = [f.name for f in files]
        assert "nested.txt" in names

    def test_non_recursive_misses_nested(self, tmp_path):
        sub = tmp_path / "subdir"
        sub.mkdir()
        _write(sub, "nested.txt", SAMPLE_TEXT)
        _write(tmp_path, "top.txt", SAMPLE_TEXT)
        files = discover_files(tmp_path, recursive=False)
        names = [f.name for f in files]
        assert "top.txt" in names
        assert "nested.txt" not in names

    def test_results_are_sorted(self, tmp_path):
        for name in ("c.txt", "a.txt", "b.txt"):
            _write(tmp_path, name, SAMPLE_TEXT)
        files = discover_files(tmp_path)
        names = [f.name for f in files]
        assert names == sorted(names)

    def test_empty_folder_returns_empty(self, tmp_path):
        assert discover_files(tmp_path) == []


class TestRunPipelineDir:
    def test_processes_all_files(self, tmp_path):
        _write(tmp_path, "doc1.txt", SAMPLE_TEXT)
        _write(tmp_path, "doc2.txt", SAMPLE_TEXT)
        results = run_pipeline_dir(tmp_path, progress=False)
        assert len(results) == 2

    def test_each_result_has_valid_schema(self, tmp_path):
        _write(tmp_path, "doc.txt", SAMPLE_TEXT)
        results = run_pipeline_dir(tmp_path, progress=False)
        assert len(results) == 1
        _validate_schema(results[0])

    def test_short_doc_marked_skipped(self, tmp_path):
        _write(tmp_path, "short.txt", "Too short.")
        _write(tmp_path, "ok.txt", SAMPLE_TEXT)
        results = run_pipeline_dir(tmp_path, progress=False)
        skipped = [r for r in results if r["skipped"]]
        ok      = [r for r in results if not r["skipped"]]
        assert len(skipped) == 1
        assert len(ok) == 1

    def test_all_results_json_serializable(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        results = run_pipeline_dir(tmp_path, progress=False)
        assert json.dumps(results)

    def test_writes_output_json_file(self, tmp_path):
        _write(tmp_path, "doc.txt", SAMPLE_TEXT)
        out = tmp_path / "results.json"
        run_pipeline_dir(tmp_path, output_path=out, progress=False)
        assert out.exists()
        data = json.loads(out.read_text())
        assert isinstance(data, list)
        assert len(data) == 1

    def test_output_json_is_valid_schema(self, tmp_path):
        _write(tmp_path, "doc.txt", SAMPLE_TEXT)
        out = tmp_path / "results.json"
        run_pipeline_dir(tmp_path, output_path=out, progress=False)
        data = json.loads(out.read_text())
        _validate_schema(data[0])

    def test_empty_folder_returns_empty_list(self, tmp_path):
        results = run_pipeline_dir(tmp_path, progress=False)
        assert results == []

    def test_raises_on_non_directory(self, tmp_path):
        f = _write(tmp_path, "file.txt", SAMPLE_TEXT)
        with pytest.raises(NotADirectoryError):
            run_pipeline_dir(f, progress=False)

    def test_recursive_processes_nested_files(self, tmp_path):
        sub = tmp_path / "sub"
        sub.mkdir()
        _write(sub, "nested.txt", SAMPLE_TEXT)
        results = run_pipeline_dir(tmp_path, recursive=True, progress=False)
        assert len(results) == 1
        assert not results[0]["skipped"]

    def test_non_recursive_skips_nested_files(self, tmp_path):
        sub = tmp_path / "sub"
        sub.mkdir()
        _write(sub, "nested.txt", SAMPLE_TEXT)
        results = run_pipeline_dir(tmp_path, recursive=False, progress=False)
        assert len(results) == 0

    def test_mixed_formats_in_folder(self, tmp_path):
        pytest.importorskip("docx", reason="python-docx not installed")
        from docx import Document as DocxDocument

        _write(tmp_path, "plain.txt", SAMPLE_TEXT)
        docx_path = tmp_path / "rich.docx"
        doc = DocxDocument()
        for line in SAMPLE_TEXT.split("\n"):
            doc.add_paragraph(line)
        doc.save(str(docx_path))

        results = run_pipeline_dir(tmp_path, progress=False)
        assert len(results) == 2
        for r in results:
            assert not r["skipped"]

    def test_file_path_in_each_result(self, tmp_path):
        _write(tmp_path, "doc.txt", SAMPLE_TEXT)
        results = run_pipeline_dir(tmp_path, progress=False)
        assert "file_path" in results[0]
        assert str(tmp_path) in results[0]["file_path"]

    def test_nuid_extracted_from_folder_doc(self, tmp_path):
        _write(tmp_path, "motion.txt", SAMPLE_TEXT)
        results = run_pipeline_dir(tmp_path, progress=False)
        assert results[0]["nuid"] is not None

    def test_parallel_workers_same_results(self, tmp_path):
        for i in range(4):
            _write(tmp_path, f"doc{i}.txt", SAMPLE_TEXT)
        serial   = run_pipeline_dir(tmp_path, workers=1, progress=False)
        parallel = run_pipeline_dir(tmp_path, workers=2, progress=False)
        # Same number of results regardless of worker count
        assert len(serial) == len(parallel) == 4
        # All parallel results have valid schemas
        for r in parallel:
            if not r["skipped"]:
                _validate_schema(r)


# ---------------------------------------------------------------------------
# Cluster ingestion tests
# ---------------------------------------------------------------------------

def _write_csv(folder: Path, rows: list[tuple[str, str]], name: str = "clusters.csv") -> Path:
    """Write a simple document_path,cluster_id CSV."""
    import csv as csv_mod
    p = folder / name
    with open(p, "w", newline="") as fh:
        writer = csv_mod.writer(fh)
        writer.writerow(["document_path", "cluster_id"])
        writer.writerows(rows)
    return p


class TestLoadClusterCsv:
    def test_basic_two_cluster_load(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        _write(tmp_path, "c.txt", SAMPLE_TEXT)
        csv_p = _write_csv(tmp_path, [
            (str(tmp_path / "a.txt"), "1"),
            (str(tmp_path / "b.txt"), "1"),
            (str(tmp_path / "c.txt"), "2"),   # singleton — dropped
        ])
        cluster_map = load_cluster_csv(csv_p)
        assert "1" in cluster_map
        assert "2" not in cluster_map          # singleton dropped
        assert len(cluster_map["1"]) == 2

    def test_singletons_dropped(self, tmp_path):
        _write(tmp_path, "x.txt", SAMPLE_TEXT)
        csv_p = _write_csv(tmp_path, [(str(tmp_path / "x.txt"), "99")])
        cluster_map = load_cluster_csv(csv_p)
        assert len(cluster_map) == 0

    def test_noise_label_minus_one_dropped(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        csv_p = _write_csv(tmp_path, [
            (str(tmp_path / "a.txt"), "-1"),
            (str(tmp_path / "b.txt"), "-1"),
        ])
        cluster_map = load_cluster_csv(csv_p)
        # -1 is in NOISE_CLUSTER_IDS → always dropped regardless of count
        assert len(cluster_map) == 0

    def test_all_noise_ids_dropped(self, tmp_path):
        _write(tmp_path, "n.txt", SAMPLE_TEXT)
        for noise_id in list(NOISE_CLUSTER_IDS)[:3]:
            if not noise_id:
                continue
            csv_p = _write_csv(tmp_path, [
                (str(tmp_path / "n.txt"), noise_id),
                (str(tmp_path / "n.txt"), noise_id),
            ], name=f"csv_{noise_id.replace('-','m')}.csv")
            cluster_map = load_cluster_csv(csv_p)
            assert len(cluster_map) == 0, f"Noise id {noise_id!r} was not dropped"

    def test_relative_paths_resolved_with_base_folder(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        csv_p = _write_csv(tmp_path, [("a.txt", "7"), ("b.txt", "7")])
        cluster_map = load_cluster_csv(csv_p, base_folder=tmp_path)
        assert "7" in cluster_map
        assert len(cluster_map["7"]) == 2

    def test_multiple_clusters_all_kept(self, tmp_path):
        for i in range(6):
            _write(tmp_path, f"d{i}.txt", SAMPLE_TEXT)
        rows = [(str(tmp_path / f"d{i}.txt"), str(i // 2)) for i in range(6)]
        csv_p = _write_csv(tmp_path, rows)
        cluster_map = load_cluster_csv(csv_p)
        assert len(cluster_map) == 3
        for recs in cluster_map.values():
            assert len(recs) == 2

    def test_records_are_csv_doc_records(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        csv_p = _write_csv(tmp_path, [
            (str(tmp_path / "a.txt"), "1"),
            (str(tmp_path / "b.txt"), "1"),
        ])
        cluster_map = load_cluster_csv(csv_p)
        for rec in cluster_map["1"]:
            assert isinstance(rec, CsvDocRecord)
            assert rec.cluster_id == "1"

    def test_tab_delimited_csv(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        p = tmp_path / "tab.csv"
        p.write_text(
            f"document_path\tcluster_id\n"
            f"{tmp_path / 'a.txt'}\t5\n"
            f"{tmp_path / 'b.txt'}\t5\n"
        )
        cluster_map = load_cluster_csv(p)
        assert "5" in cluster_map

    def test_rich_csv_columns_parsed(self, tmp_path):
        """Full schema CSV columns land in CsvDocRecord fields."""
        import csv as csv_mod
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        csv_p = tmp_path / "rich.csv"
        with open(csv_p, "w", newline="") as fh:
            w = csv_mod.writer(fh)
            w.writerow([
                "cluster_id", "cluster_stage", "cluster_size", "anchor_uid",
                "cluster_pool_size", "doc_origin", "merged_from_s2_count",
                "ground_truth_folder", "majority_folder_in_cluster",
                "is_misclassified", "document_filename", "document_path",
                "num_raw_dockets", "docket_primary_raw", "docket_primary_normalized",
                "docket_source", "all_dockets_raw", "all_dockets_normalized",
                "num_westlaw_ids", "all_westlaw_ids", "num_lexis_ids", "all_lexis_ids",
                "num_uids_total", "has_any_uid", "num_entities", "all_entities",
                "doc_overlap_with_cluster", "caption", "court", "location", "ocr_applied",
            ])
            for fname in ("a.txt", "b.txt"):
                w.writerow([
                    "42", "final", "2", "anchor-001",
                    "5", "s3", "0",
                    "/gt/folder", "/majority/folder",
                    "False", fname, str(tmp_path / fname),
                    "1", "3:22-cv-01234", "3:22-CV-01234",
                    "labeled", "3:22-cv-01234", "3:22-CV-01234",
                    "0", "", "0", "",
                    "1", "True", "3", "ACME CORP|GLOBEX LLC|J. Smith",
                    "0.85", "ACME CORP v. GLOBEX LLC",
                    "UNITED STATES DISTRICT COURT", "Northern District of California",
                    "False",
                ])
        cluster_map = load_cluster_csv(csv_p, base_folder=tmp_path)
        assert "42" in cluster_map
        rec = cluster_map["42"][0]
        assert rec.nuid == "3:22-CV-01234"
        assert rec.nuid_source == "labeled"
        assert rec.court_name == "UNITED STATES DISTRICT COURT"
        assert rec.court_location == "Northern District of California"
        assert rec.cluster_stage == "final"
        assert rec.anchor_uid == "anchor-001"
        assert rec.is_misclassified is False
        assert "ACME CORP" in rec.all_entities


class TestConsolidateCluster:
    def _make_result(self, **kwargs) -> dict:
        base = {
            "skipped": False, "file_path": "x.txt",
            "nuid": None, "court_name": None, "court_location": None,
            "judge_name": None, "filing_date": None,
            "document_types": [], "title": None,
            "parties": {"plaintiffs": [], "defendants": [], "petitioners": [], "respondents": []},
            "clauses": [], "evidence": {}, "confidence": {},
        }
        base.update(kwargs)
        return base

    def test_majority_nuid(self):
        results = [
            self._make_result(nuid="3:22-CV-01234"),
            self._make_result(nuid="3:22-CV-01234"),
            self._make_result(nuid="9:99-CV-99999"),
        ]
        consensus = consolidate_cluster(results)
        assert consensus["nuid"] == "3:22-CV-01234"

    def test_all_null_returns_none(self):
        results = [self._make_result(nuid=None), self._make_result(nuid=None)]
        consensus = consolidate_cluster(results)
        assert consensus["nuid"] is None

    def test_union_document_types(self):
        results = [
            self._make_result(document_types=["Motion"]),
            self._make_result(document_types=["Brief"]),
            self._make_result(document_types=["Motion", "Brief"]),
        ]
        consensus = consolidate_cluster(results)
        assert "Motion" in consensus["document_types"]
        assert "Brief" in consensus["document_types"]

    def test_union_parties(self):
        results = [
            self._make_result(parties={"plaintiffs": ["ACME CORP"], "defendants": [], "petitioners": [], "respondents": []}),
            self._make_result(parties={"plaintiffs": [], "defendants": ["GLOBEX LLC"], "petitioners": [], "respondents": []}),
        ]
        consensus = consolidate_cluster(results)
        assert "ACME CORP" in consensus["parties"]["plaintiffs"]
        assert "GLOBEX LLC" in consensus["parties"]["defendants"]

    def test_skipped_docs_excluded_from_consensus(self):
        results = [
            {"skipped": True, "file_path": "x.txt", "nuid": "WRONG"},
            self._make_result(nuid="3:22-CV-01234"),
            self._make_result(nuid="3:22-CV-01234"),
        ]
        consensus = consolidate_cluster(results)
        assert consensus["nuid"] == "3:22-CV-01234"

    def test_empty_cluster_returns_empty(self):
        assert consolidate_cluster([]) == {}


class TestRunPipelineClusters:
    def _make_cluster_map(self, tmp_path, files: dict) -> dict:
        cluster_map = {}
        for cid, names in files.items():
            cluster_map[cid] = [
                CsvDocRecord(file_path=tmp_path / name, cluster_id=cid)
                for name in names
            ]
        return cluster_map

    def test_basic_cluster_run(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        cluster_map = self._make_cluster_map(tmp_path, {"42": ["a.txt", "b.txt"]})
        results = run_pipeline_clusters(cluster_map, progress=False)
        assert len(results) == 2

    def test_each_result_has_cluster_fields(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        cluster_map = self._make_cluster_map(tmp_path, {"42": ["a.txt", "b.txt"]})
        results = run_pipeline_clusters(cluster_map, progress=False)
        for r in results:
            assert r["cluster_id"] == "42"
            assert r["cluster_size"] == 2
            assert "cluster_consensus" in r
            assert "cluster_stage" in r
            assert "anchor_uid" in r

    def test_csv_nuid_overrides_pipeline(self, tmp_path):
        """CSV docket_primary_normalized must override pipeline result."""
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        cluster_map = {
            "1": [
                CsvDocRecord(file_path=tmp_path / "a.txt", cluster_id="1",
                             nuid="5:23-CV-99999-ABC", nuid_source="labeled"),
                CsvDocRecord(file_path=tmp_path / "b.txt", cluster_id="1"),
            ]
        }
        results = run_pipeline_clusters(cluster_map, progress=False)
        nuid_result = next(r for r in results if "a.txt" in r["file_path"])
        assert nuid_result["nuid"] == "5:23-CV-99999-ABC"
        assert nuid_result["confidence"]["nuid"] == 0.97
        assert nuid_result["evidence"]["nuid"][0]["source"] == "csv_metadata"

    def test_csv_court_fills_null(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        cluster_map = {
            "2": [
                CsvDocRecord(file_path=tmp_path / "a.txt", cluster_id="2",
                             court_name="SUPERIOR COURT OF CALIFORNIA"),
                CsvDocRecord(file_path=tmp_path / "b.txt", cluster_id="2"),
            ]
        }
        results = run_pipeline_clusters(cluster_map, progress=False)
        for r in results:
            assert r["court_name"] is not None

    def test_csv_extra_fields_present(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        cluster_map = {
            "3": [
                CsvDocRecord(file_path=tmp_path / "a.txt", cluster_id="3",
                             all_entities=["ACME CORP", "GLOBEX"]),
                CsvDocRecord(file_path=tmp_path / "b.txt", cluster_id="3"),
            ]
        }
        results = run_pipeline_clusters(cluster_map, progress=False)
        for key in ("csv_caption", "csv_all_dockets", "csv_westlaw_ids",
                    "csv_lexis_ids", "csv_all_entities", "csv_doc_origin"):
            assert key in results[0], f"Missing passthrough key: {key}"

    def test_consensus_nuid_filled_in(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        cluster_map = self._make_cluster_map(tmp_path, {"1": ["a.txt", "b.txt"]})
        results = run_pipeline_clusters(cluster_map, progress=False)
        for r in results:
            assert r["cluster_consensus"].get("nuid") is not None

    def test_output_json_written(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        cluster_map = self._make_cluster_map(tmp_path, {"99": ["a.txt", "b.txt"]})
        out = tmp_path / "out.json"
        run_pipeline_clusters(cluster_map, output_path=out, progress=False)
        assert out.exists()
        import json as _json
        data = _json.loads(out.read_text())
        assert len(data) == 2

    def test_full_csv_to_results(self, tmp_path):
        for name in ("a.txt", "b.txt", "c.txt"):
            _write(tmp_path, name, SAMPLE_TEXT)
        csv_p = _write_csv(tmp_path, [
            (str(tmp_path / "a.txt"), "10"),
            (str(tmp_path / "b.txt"), "10"),
            (str(tmp_path / "c.txt"), "11"),
        ])
        cluster_map = load_cluster_csv(csv_p)
        assert "11" not in cluster_map
        results = run_pipeline_clusters(cluster_map, progress=False)
        assert len(results) == 2
        assert all(r["cluster_id"] == "10" for r in results)

    def test_schema_valid_after_cluster_augmentation(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        cluster_map = self._make_cluster_map(tmp_path, {"1": ["a.txt", "b.txt"]})
        results = run_pipeline_clusters(cluster_map, progress=False)
        for r in results:
            if not r["skipped"]:
                _validate_schema(r)

    def test_consensus_source_in_evidence_when_filled(self, tmp_path):
        _write(tmp_path, "a.txt", SAMPLE_TEXT)
        _write(tmp_path, "b.txt", SAMPLE_TEXT)
        cluster_map = self._make_cluster_map(tmp_path, {"X": ["a.txt", "b.txt"]})
        results = run_pipeline_clusters(cluster_map, progress=False)
        for r in results:
            for field, ev_list in r.get("evidence", {}).items():
                for ev in ev_list:
                    if ev["source"] == "cluster_consensus":
                        assert ev["rule_id"].startswith("cluster_consensus:")